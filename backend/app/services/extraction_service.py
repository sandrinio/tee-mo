"""
Document extraction service — EPIC-014, STORY-014-01.

Provides format-specific text extraction functions for PDF, DOCX, and XLSX
files, plus a shared truncation helper.  These functions are intentionally
free of Google Drive dependencies so that the local-upload endpoint
(STORY-014-02) can reuse them without importing the Drive client.

Public API:
  - extract_pdf(raw_bytes: bytes) -> str
  - extract_docx(raw_bytes: bytes) -> str
  - extract_xlsx(raw_bytes: bytes) -> str
  - maybe_truncate(content: str) -> str
  - rows_to_markdown_table(rows: list[tuple]) -> str
  - docx_table_to_markdown(table) -> str
"""

from __future__ import annotations

import io

# Module-level imports required so tests can monkeypatch these names.
import pymupdf4llm
import pymupdf
from docx import Document as DocxDocument
from openpyxl import load_workbook

# ---------------------------------------------------------------------------
# Content truncation threshold (ADR-016)
# ---------------------------------------------------------------------------

_TRUNCATION_LIMIT = 50_000
_TRUNCATION_NOTICE = "\n\n[Content truncated at 50000 characters]"


def rows_to_markdown_table(rows: list[tuple]) -> str:
    """Convert a list of row tuples into a pipe-separated markdown table.

    First row becomes the header. None values become empty strings.
    """
    if not rows:
        return ""

    def cell(v):
        return str(v) if v is not None else ""

    header = rows[0]
    col_count = len(header)
    lines = [
        "| " + " | ".join(cell(v) for v in header) + " |",
        "| " + " | ".join("---" for _ in range(col_count)) + " |",
    ]
    for row in rows[1:]:
        cells = [cell(row[i]) if i < len(row) else "" for i in range(col_count)]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def docx_table_to_markdown(table) -> str:
    """Convert a python-docx Table object to a markdown table."""
    rows = []
    for row in table.rows:
        rows.append(tuple(cell.text.strip() for cell in row.cells))
    return rows_to_markdown_table(rows)


def extract_pdf(raw_bytes: bytes) -> str:
    """Extract text from a PDF as markdown using pymupdf4llm.

    Args:
        raw_bytes: Raw PDF file bytes.

    Returns:
        Markdown-formatted text preserving tables, headings, and layout.
    """
    doc = pymupdf.open(stream=raw_bytes, filetype="pdf")
    md = pymupdf4llm.to_markdown(doc)
    doc.close()
    return md


def extract_docx(raw_bytes: bytes) -> str:
    """Extract text + tables from a DOCX file, preserving document order.

    Args:
        raw_bytes: Raw DOCX file bytes.

    Returns:
        Text with tables rendered as markdown tables.
    """
    doc = DocxDocument(io.BytesIO(raw_bytes))
    parts = []
    for element in doc.element.body:
        tag = element.tag.split("}")[-1]
        if tag == "p":
            for para in doc.paragraphs:
                if para._element is element:
                    if para.text.strip():
                        parts.append(para.text)
                    break
        elif tag == "tbl":
            for table in doc.tables:
                if table._element is element:
                    parts.append(docx_table_to_markdown(table))
                    break
    return "\n\n".join(parts)


def extract_xlsx(raw_bytes: bytes) -> str:
    """Extract all sheets from an XLSX as named markdown tables.

    Args:
        raw_bytes: Raw XLSX file bytes.

    Returns:
        Markdown with ## Sheet: headers and pipe-separated tables per sheet.
    """
    wb = load_workbook(io.BytesIO(raw_bytes))
    sheet_texts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        md = f"## Sheet: {sheet_name}\n\n{rows_to_markdown_table(rows)}"
        sheet_texts.append(md)
    return "\n\n".join(sheet_texts)


def maybe_truncate(content: str) -> str:
    """Truncate content at 50,000 characters and append a trim notice if needed.

    Args:
        content: Extracted text string (may be any length).

    Returns:
        Original string if <= 50,000 chars, or first 50,000 chars + trim notice.
    """
    if len(content) > _TRUNCATION_LIMIT:
        return content[:_TRUNCATION_LIMIT] + _TRUNCATION_NOTICE
    return content
