"""
Google Drive content extraction service — EPIC-006, STORY-006-01/006-07.

Provides three public functions:
  - get_drive_client: decrypt refresh token and build an authenticated Drive API client.
  - fetch_file_content: extract text from a Drive file by MIME type.
  - compute_content_hash: MD5 digest of content string for change detection.

MIME type support (ADR-016, updated STORY-006-07):
  - application/vnd.google-apps.document      → export as text/plain
  - application/vnd.google-apps.spreadsheet   → export as XLSX → _extract_xlsx (markdown tables)
  - application/vnd.google-apps.presentation  → export as text/plain → _add_slide_markers
  - application/pdf                            → get_media + pymupdf4llm.to_markdown
  - application/vnd.openxmlformats-officedocument.wordprocessingml.document → get_media + python-docx (with tables)
  - application/vnd.openxmlformats-officedocument.spreadsheetml.sheet       → get_media + openpyxl (markdown tables)

All imports that tests monkeypatch are at module level so monkeypatch.setattr works.
"""

from __future__ import annotations

import asyncio
import hashlib
import io

# Module-level imports required so tests can monkeypatch these names.
# (See FLASHCARDS.md httpx module-level import rule — same pattern applies here.)
from app.core.encryption import decrypt
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import pymupdf4llm
import pymupdf
from docx import Document as DocxDocument
from openpyxl import load_workbook

from app.core.config import get_settings

# ---------------------------------------------------------------------------
# Content truncation threshold (ADR-016)
# ---------------------------------------------------------------------------

_TRUNCATION_LIMIT = 50_000
_TRUNCATION_NOTICE = "\n\n[Content truncated at 50000 characters]"

# ---------------------------------------------------------------------------
# Multimodal fallback constants (STORY-006-08)
#
# PDFs returning fewer than _MULTIMODAL_FALLBACK_THRESHOLD characters from
# pymupdf4llm are treated as scanned/image-only documents.  For supported
# providers (google, openai) the file is re-processed via extract_content_multimodal.
# Anthropic cannot process raw PDFs — a user-visible warning is appended instead.
# Files larger than _MULTIMODAL_SIZE_LIMIT skip the AI fallback regardless of provider.
# ---------------------------------------------------------------------------

_SCANNED_PDF_WARNING = (
    "\n\n[Warning: This file appears to be a scanned document. "
    "Text extraction is limited for this provider. Consider using "
    "Google or OpenAI as your AI provider for better scanned document support.]"
)
_OVERSIZED_PDF_WARNING = (
    "\n\n[Warning: This scanned document is too large for "
    "AI-assisted extraction. Text shown is best-effort.]"
)
_MULTIMODAL_FALLBACK_THRESHOLD = 100
_MULTIMODAL_SIZE_LIMIT = 20_000_000  # 20 MB


class _AwaitableStr(str):
    """A str subclass that is also awaitable.

    Returned by fetch_file_content so that the function can be called both
    synchronously (legacy callers without provider/api_key) and with ``await``
    (new async callers added in STORY-006-08) without changing the function
    signature from a regular ``def`` to ``async def`` — which would break the
    existing synchronous test suite.

    When used as a plain string (e.g., ``"foo" in result``), behaviour is
    identical to a normal ``str``.  When awaited, a completed asyncio ``Future``
    is returned so the event loop yields once and then produces the plain string.
    """

    def __await__(self):
        """Return an iterator over a completed Future holding this string's value."""
        f = asyncio.get_event_loop().create_future()
        f.set_result(str(self))
        return f.__await__()

# ---------------------------------------------------------------------------
# MIME type routing
# ---------------------------------------------------------------------------

# Google Workspace export targets (ADR-016, updated STORY-006-07)
# Sheets now export as XLSX so _extract_xlsx can produce markdown tables.
_GOOGLE_EXPORT_MIMES: dict[str, str] = {
    "application/vnd.google-apps.document": "text/plain",
    "application/vnd.google-apps.spreadsheet": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.google-apps.presentation": "text/plain",
}

# Binary file MIME types handled via get_media + parser
_BINARY_MIMES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}

# Plain-text MIME types downloaded directly as UTF-8 text
_TEXT_MIMES = {
    "text/plain",
}


def get_drive_client(encrypted_refresh_token: str):
    """Decrypt an encrypted OAuth refresh token and build an authenticated Drive v3 client.

    Steps:
      1. Decrypt the refresh token using AES-256-GCM (app.core.encryption.decrypt).
      2. Construct a google.oauth2.credentials.Credentials instance with the
         decrypted refresh token and OAuth app credentials from Settings.
      3. Call creds.refresh(Request()) to exchange the refresh token for a
         short-lived access token.
      4. Build and return a googleapiclient Drive v3 service object.

    Args:
        encrypted_refresh_token: AES-256-GCM encrypted refresh token ciphertext,
            as stored in the database. NEVER log the plaintext.

    Returns:
        A googleapiclient Resource object for the Drive v3 API.
    """
    settings = get_settings()
    plaintext_token = decrypt(encrypted_refresh_token)

    creds = Credentials(
        token=None,
        refresh_token=plaintext_token,
        client_id=settings.google_api_client_id,
        client_secret=settings.google_api_secret,
        token_uri="https://oauth2.googleapis.com/token",
    )
    creds.refresh(Request())

    return build("drive", "v3", credentials=creds)


def _download_media(drive_client, drive_file_id: str) -> bytes:
    """Download binary file content from Drive using MediaIoBaseDownload.

    Args:
        drive_client: Authenticated Drive v3 API client.
        drive_file_id: Google Drive file ID.

    Returns:
        Raw bytes of the downloaded file.
    """
    request = drive_client.files().get_media(fileId=drive_file_id)
    buffer = io.BytesIO()
    downloader = MediaIoBaseDownload(buffer, request)
    done = False
    while not done:
        _status, done = downloader.next_chunk()
    buffer.seek(0)
    return buffer.read()


def _rows_to_markdown_table(rows: list[tuple]) -> str:
    """Convert a list of row tuples into a pipe-separated markdown table.

    First row becomes the header. None values become empty strings.
    """
    if not rows:
        return ""

    def cell(v):
        return str(v) if v is not None else ""

    header = rows[0]
    col_count = len(header)
    lines = [
        "| " + " | ".join(cell(v) for v in header) + " |",
        "| " + " | ".join("---" for _ in range(col_count)) + " |",
    ]
    for row in rows[1:]:
        cells = [cell(row[i]) if i < len(row) else "" for i in range(col_count)]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def _docx_table_to_markdown(table) -> str:
    """Convert a python-docx Table object to a markdown table."""
    rows = []
    for row in table.rows:
        rows.append(tuple(cell.text.strip() for cell in row.cells))
    return _rows_to_markdown_table(rows)


def _extract_pdf(raw_bytes: bytes) -> str:
    """Extract text from a PDF as markdown using pymupdf4llm.

    Args:
        raw_bytes: Raw PDF file bytes.

    Returns:
        Markdown-formatted text preserving tables, headings, and layout.
    """
    doc = pymupdf.open(stream=raw_bytes, filetype="pdf")
    md = pymupdf4llm.to_markdown(doc)
    doc.close()
    return md


def _extract_docx(raw_bytes: bytes) -> str:
    """Extract text + tables from a DOCX file, preserving document order.

    Args:
        raw_bytes: Raw DOCX file bytes.

    Returns:
        Text with tables rendered as markdown tables.
    """
    doc = DocxDocument(io.BytesIO(raw_bytes))
    parts = []
    for element in doc.element.body:
        tag = element.tag.split("}")[-1]
        if tag == "p":
            for para in doc.paragraphs:
                if para._element is element:
                    if para.text.strip():
                        parts.append(para.text)
                    break
        elif tag == "tbl":
            for table in doc.tables:
                if table._element is element:
                    parts.append(_docx_table_to_markdown(table))
                    break
    return "\n\n".join(parts)


def _extract_xlsx(raw_bytes: bytes) -> str:
    """Extract all sheets from an XLSX as named markdown tables.

    Args:
        raw_bytes: Raw XLSX file bytes.

    Returns:
        Markdown with ## Sheet: headers and pipe-separated tables per sheet.
    """
    wb = load_workbook(io.BytesIO(raw_bytes))
    sheet_texts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        md = f"## Sheet: {sheet_name}\n\n{_rows_to_markdown_table(rows)}"
        sheet_texts.append(md)
    return "\n\n".join(sheet_texts)


def _add_slide_markers(text: str) -> str:
    """Insert slide boundary markers into Slides plain text export.

    Google Slides text/plain export separates slides with form-feed (\\x0c).
    Falls back to double-newline splitting if no form-feeds found.
    """
    parts = text.split("\x0c")
    if len(parts) <= 1:
        parts = [p.strip() for p in text.split("\n\n") if p.strip()]
    if len(parts) <= 1:
        return text
    return "\n\n".join(
        f"--- Slide {i + 1} ---\n{part.strip()}" for i, part in enumerate(parts)
    )


def _maybe_truncate(content: str) -> str:
    """Truncate content at 50,000 characters and append a trim notice if needed.

    Args:
        content: Extracted text string (may be any length).

    Returns:
        Original string if <= 50,000 chars, or first 50,000 chars + trim notice.
    """
    if len(content) > _TRUNCATION_LIMIT:
        return content[:_TRUNCATION_LIMIT] + _TRUNCATION_NOTICE
    return content


def fetch_file_content(
    drive_client,
    drive_file_id: str,
    mime_type: str,
    *,
    provider: str | None = None,
    api_key: str | None = None,
) -> "_AwaitableStr":
    """Extract text content from a Google Drive file by MIME type.

    Dispatches to the appropriate extraction strategy based on mime_type:
      - Google Docs: export as text/plain.
      - Google Sheets: export as XLSX, then extract as markdown tables.
      - Google Slides: export as text/plain, then add slide boundary markers.
      - PDF: downloads binary via get_media, extracts markdown with pymupdf4llm.
        If the extracted text is fewer than _MULTIMODAL_FALLBACK_THRESHOLD characters
        (indicating a scanned/image-only PDF), a multimodal fallback is attempted:
          - google / openai + file ≤ 20 MB  → delegates to scan_service.extract_content_multimodal
          - google / openai + file >  20 MB → appends _OVERSIZED_PDF_WARNING
          - anthropic                        → appends _SCANNED_PDF_WARNING
        Falls back gracefully (no warning) when provider/api_key are not supplied.
      - DOCX: downloads binary via get_media, extracts text + tables with python-docx.
      - XLSX: downloads binary via get_media, extracts markdown tables with openpyxl.
      - Unsupported types: raises ValueError.

    Content is truncated at 50,000 characters with a trim notice appended
    (ADR-016 content cap to prevent LLM context overflow).

    The return type is ``_AwaitableStr`` — a ``str`` subclass that also implements
    ``__await__``.  This allows the function to be called both synchronously
    (legacy callers without provider/api_key) and with ``await`` (new STORY-006-08
    async callers) without changing the function to ``async def``.

    When the multimodal fallback IS needed (provider in ("google", "openai") and
    file is small enough), this function returns a *coroutine* object directly.
    The caller must ``await`` it in that case — which is always true for the
    new async callers that supply provider/api_key.

    Args:
        drive_client: Authenticated Drive v3 API client (from get_drive_client).
        drive_file_id: Google Drive file ID.
        mime_type: MIME type of the file (as returned by the Drive files.list API).
        provider: Optional BYOK provider slug — "google", "anthropic", or "openai".
            Required to enable the multimodal fallback for scanned PDFs.
        api_key: Optional decrypted BYOK API key.  Must be supplied together with provider.

    Returns:
        An ``_AwaitableStr`` containing the extracted text (truncated at 50 000 chars),
        or a coroutine that resolves to such a string when the multimodal fallback path
        is taken.

    Raises:
        ValueError: If mime_type is not one of the 6 supported types (ADR-016).
    """
    if mime_type in _GOOGLE_EXPORT_MIMES:
        export_mime = _GOOGLE_EXPORT_MIMES[mime_type]
        raw = drive_client.files().export(fileId=drive_file_id, mimeType=export_mime).execute()
        if export_mime.endswith("spreadsheetml.sheet"):
            content = _extract_xlsx(raw if isinstance(raw, bytes) else raw.encode())
        elif mime_type == "application/vnd.google-apps.presentation":
            text = raw.decode("utf-8") if isinstance(raw, bytes) else raw
            content = _add_slide_markers(text)
        else:
            content = raw.decode("utf-8") if isinstance(raw, bytes) else raw

    elif mime_type == "application/pdf":
        raw_bytes = _download_media(drive_client, drive_file_id)
        content = _extract_pdf(raw_bytes)

        # Multimodal fallback for scanned/image-based PDFs (STORY-006-08).
        # Only triggered when pymupdf4llm returns fewer than the threshold chars
        # AND both provider and api_key are supplied by the caller.
        if len(content.strip()) < _MULTIMODAL_FALLBACK_THRESHOLD and provider and api_key:
            if len(raw_bytes) > _MULTIMODAL_SIZE_LIMIT:
                # File is too large for AI-assisted extraction regardless of provider.
                content += _OVERSIZED_PDF_WARNING
            elif provider in ("google", "openai"):
                # Return a coroutine — the async caller will await it.
                # The lazy import ensures patch() context managers in tests work correctly.
                async def _do_fallback(
                    _raw_bytes=raw_bytes, _provider=provider, _api_key=api_key
                ) -> str:
                    """Run scan-tier multimodal extraction and apply content truncation."""
                    from app.services.scan_service import extract_content_multimodal
                    fb_content = await extract_content_multimodal(
                        _raw_bytes, _provider, _api_key
                    )
                    return _maybe_truncate(fb_content)

                return _do_fallback()
            elif provider == "anthropic":
                # Anthropic cannot process raw PDFs natively — warn the user instead.
                content += _SCANNED_PDF_WARNING

    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        raw_bytes = _download_media(drive_client, drive_file_id)
        content = _extract_docx(raw_bytes)

    elif mime_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        raw_bytes = _download_media(drive_client, drive_file_id)
        content = _extract_xlsx(raw_bytes)

    elif mime_type in _TEXT_MIMES:
        # Plain text files: download via get_media and decode as UTF-8.
        raw_bytes = _download_media(drive_client, drive_file_id)
        content = raw_bytes.decode("utf-8", errors="replace")

    else:
        raise ValueError(
            f"Unsupported MIME type: {mime_type!r}. "
            f"Supported types: {sorted(_GOOGLE_EXPORT_MIMES.keys() | _BINARY_MIMES | _TEXT_MIMES)}"
        )

    return _AwaitableStr(_maybe_truncate(content))


def compute_content_hash(content: str) -> str:
    """Compute an MD5 hex digest of the content string for change detection.

    Used by the self-healing description mechanism (ADR-006): if the hash of
    the current Drive file content differs from the stored hash, the
    ai_description is re-generated.

    Args:
        content: Text content string to hash.

    Returns:
        32-character lowercase hexadecimal MD5 digest.
    """
    return hashlib.md5(content.encode()).hexdigest()
